# -*- coding: utf-8 -*-
"""
自研 vs EASE5 直达声压对比测试报告生成脚本
按场景+声源维度汇总，生成Word报告，含散点图（±3dB红框标注）
"""

import os
import re
import warnings
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import rcParams
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from openpyxl import load_workbook

warnings.filterwarnings('ignore')

# 设置中文字体
rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
rcParams['axes.unicode_minus'] = False

BASE_DIR = r'e:\AudioSimulationSnap\点位声压采样\求解结果对比-直达声压（RoomMapping）'
OUTPUT_DIR = r'e:\AudioSimulationSnap\点位声压采样'
REPORT_FILE = os.path.join(OUTPUT_DIR, '自研vs_EASE5_直达声压对比测试报告.docx')
CHART_DIR = os.path.join(OUTPUT_DIR, 'charts')
os.makedirs(CHART_DIR, exist_ok=True)

# 频点顺序
FREQ_ORDER = [125, 250, 500, 1000, 2000, 4000, 8000]

# 场景显示名映射
SCENE_MAP = {
    '中型教室1': '中型教室1',
    '会议室': '会议室',
    '超大型教室': '超大型教室',
}

# 声源模式映射
SOURCE_MAP = {
    '单声源': '单声源',
    '双声源': '双声源',
}


def find_excel_files():
    """查找所有对比Excel文件，排除大型教室（尺寸不匹配）和临时文件"""
    files = []
    for root, dirs, filenames in os.walk(BASE_DIR):
        # 排除"大型教室 - 由于EASE5和自研尺寸对不上暂时不测"
        rel_path = os.path.relpath(root, BASE_DIR)
        if '大型教室 - 由于EASE5和自研尺寸对不上暂时不测' in rel_path:
            continue
        for fn in filenames:
            if fn.startswith('~$'):
                continue
            if 'ES5和声压数据相对最近点位声压级对比' in fn and fn.endswith('.xlsx'):
                files.append(os.path.join(root, fn))
    return files


def parse_file_info(filepath):
    """从文件路径和文件名解析场景、声源模式、频点"""
    rel = os.path.relpath(filepath, BASE_DIR)
    parts = rel.split(os.sep)

    # 场景：第一级目录
    scene = parts[0] if len(parts) > 0 else ''
    # 声源模式：第二级目录
    source_mode = parts[1] if len(parts) > 1 else ''

    # 频点：从文件名提取
    freq_match = re.search(r'【(\d+)HZ】', os.path.basename(filepath))
    freq = int(freq_match.group(1)) if freq_match else 0

    return scene, source_mode, freq


def read_excel_data(filepath):
    """读取Excel文件的matched_points和stats sheet"""
    try:
        wb = load_workbook(filepath, data_only=True)
    except Exception as e:
        print(f"  [WARN] Cannot open {filepath}: {e}")
        return None, None, None

    # 读取数据
    ws_data = wb['matched_points']
    rows = list(ws_data.iter_rows(values_only=True))
    if len(rows) < 2:
        wb.close()
        return None, None, None

    headers = rows[0]
    data_rows = rows[1:]

    # 标准化列名（处理编码问题）
    col_map = {}
    for i, h in enumerate(headers):
        if h is None:
            continue
        h_str = str(h)
        if 'ease5_x' == h_str or h_str == 'ease5_x':
            col_map['ease5_x'] = i
        elif 'ease5_y' == h_str:
            col_map['ease5_y'] = i
        elif 'ease5_z' == h_str:
            col_map['ease5_z'] = i
        elif 'self_x' == h_str and 'ease5' not in h_str:
            col_map['self_x'] = i
        elif 'self_y' == h_str and 'ease5' not in h_str:
            col_map['self_y'] = i
        elif 'self_z' == h_str and 'ease5' not in h_str:
            col_map['self_z'] = i
        elif 'ease5_spl' in h_str:
            col_map['ease5_spl'] = i
        elif 'self_spl' in h_str and 'ease5' not in h_str:
            col_map['self_spl'] = i
        elif 'delta_db' in h_str:
            col_map['delta_db'] = i
        elif 'match_dist_xyz' in h_str:
            col_map['match_dist_xyz'] = i

    # 提取声源坐标信息
    self_source = None
    ease5_source = None
    # 声源坐标在最后两列
    for row in data_rows:
        if len(row) >= 19 and row[17] is not None:
            self_source = str(row[17])
        if len(row) >= 19 and row[18] is not None:
            ease5_source = str(row[18])
        if self_source or ease5_source:
            break

    # 提取序号列
    col_map['ease5_seq'] = 0  # ease5序号
    col_map['self_seq'] = 6   # self序号（第7列）
    # 行号列号
    col_map['ease5_row'] = 1  # ease5行号
    col_map['ease5_col'] = 2  # ease5列号
    col_map['self_row'] = 7   # self行号
    col_map['self_col'] = 8   # self列号

    # 构建DataFrame
    records = []
    for row in data_rows:
        try:
            rec = {}
            for key, idx in col_map.items():
                if idx < len(row):
                    rec[key] = row[idx]
                else:
                    rec[key] = None
            records.append(rec)
        except Exception:
            continue

    df = pd.DataFrame(records)

    # 读取stats
    stats = None
    if 'stats' in wb.sheetnames:
        ws_stats = wb['stats']
        stats_rows = list(ws_stats.iter_rows(values_only=True))
        if len(stats_rows) >= 2:
            stats = dict(zip(stats_rows[0], stats_rows[1]))

    wb.close()
    return df, stats, (self_source, ease5_source)


def collect_all_data():
    """收集所有Excel数据"""
    files = find_excel_files()
    print(f"找到 {len(files)} 个对比文件")

    all_data = []
    for fp in files:
        scene, source_mode, freq = parse_file_info(fp)
        if scene not in SCENE_MAP:
            continue
        if source_mode not in SOURCE_MAP:
            continue
        if freq == 0:
            continue

        print(f"  读取: {SCENE_MAP.get(scene, scene)} / {SOURCE_MAP.get(source_mode, source_mode)} / {freq}Hz")
        df, stats, sources = read_excel_data(fp)
        if df is None or len(df) == 0:
            print(f"    [WARN] 无数据，跳过")
            continue

        all_data.append({
            'scene': scene,
            'scene_display': SCENE_MAP.get(scene, scene),
            'source_mode': source_mode,
            'source_display': SOURCE_MAP.get(source_mode, source_mode),
            'freq': freq,
            'df': df,
            'stats': stats,
            'sources': sources,
            'filepath': fp,
        })

    return all_data


def compute_statistics(df):
    """计算delta_db的统计指标"""
    delta = df['delta_db'].dropna().astype(float)
    total = len(delta)
    if total == 0:
        return None

    within_3 = ((delta >= -3) & (delta <= 3)).sum()
    below_3 = (delta < -3).sum()
    above_3 = (delta > 3).sum()

    stats = {
        'total': total,
        'mean': delta.mean(),
        'std': delta.std(),
        'min': delta.min(),
        'max': delta.max(),
        'mae': delta.abs().mean(),
        'rmse': np.sqrt((delta**2).mean()),
        'within_3_count': int(within_3),
        'within_3_pct': within_3 / total * 100,
        'below_3_count': int(below_3),
        'below_3_pct': below_3 / total * 100,
        'above_3_count': int(above_3),
        'above_3_pct': above_3 / total * 100,
        'p95_abs': delta.abs().quantile(0.95),
        'median': delta.median(),
    }
    return stats


def generate_scatter_plot(data_item, output_path):
    """生成散点图，带±3dB红框标注"""
    df = data_item['df']
    delta = df['delta_db'].dropna().astype(float).values
    ease5_spl = df['ease5_spl'].dropna().astype(float).values

    # 使用序号作为x轴
    x = np.arange(1, len(delta) + 1)

    fig, ax = plt.subplots(figsize=(12, 5))

    # 根据偏差值着色
    colors = []
    for d in delta:
        if d > 3:
            colors.append('#FF4444')  # 红：>3
        elif d < -3:
            colors.append('#4444FF')  # 蓝：<-3
        else:
            colors.append('#44BB44')  # 绿：±3内

    ax.scatter(x, delta, c=colors, s=20, alpha=0.7, edgecolors='none')

    # ±3dB红框
    rect = mpatches.FancyBboxPatch(
        (0, -3), len(delta) + 1, 6,
        boxstyle="round,pad=0",
        linewidth=2.5, edgecolor='red', facecolor='none',
        linestyle='--', zorder=5
    )
    ax.add_patch(rect)

    # ±3dB参考线
    ax.axhline(y=3, color='red', linestyle='--', linewidth=1, alpha=0.7, label='+3 dB')
    ax.axhline(y=-3, color='red', linestyle='--', linewidth=1, alpha=0.7, label='-3 dB')
    ax.axhline(y=0, color='gray', linestyle='-', linewidth=0.8, alpha=0.5)

    # 标注±3dB区域文字
    ax.text(len(delta) + 0.5, 0, '±3dB\n容差带', fontsize=9, color='red',
            va='center', ha='left', fontweight='bold')

    stats = compute_statistics(df)
    title = f"{data_item['scene_display']} - {data_item['source_display']} - {data_item['freq']}Hz"
    ax.set_title(f'{title}\nΔdB(self-EASE5) 散点分布 (±3dB内占比: {stats["within_3_pct"]:.1f}%)',
                 fontsize=12, fontweight='bold')
    ax.set_xlabel('测点序号', fontsize=10)
    ax.set_ylabel('ΔdB (self - EASE5)', fontsize=10)

    # 图例
    from matplotlib.lines import Line2D
    legend_elements = [
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#44BB44', markersize=8, label=f'±3dB内 ({stats["within_3_count"]}个, {stats["within_3_pct"]:.1f}%)'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#FF4444', markersize=8, label=f'>+3dB ({stats["above_3_count"]}个, {stats["above_3_pct"]:.1f}%)'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#4444FF', markersize=8, label=f'<-3dB ({stats["below_3_count"]}个, {stats["below_3_pct"]:.1f}%)'),
        Line2D([0], [0], color='red', linestyle='--', linewidth=2, label='±3dB容差线'),
    ]
    ax.legend(handles=legend_elements, loc='upper right', fontsize=8)

    ax.set_xlim(0, len(delta) + 1)
    y_range = max(abs(delta.min()), abs(delta.max()), 4)
    ax.set_ylim(-y_range - 1, y_range + 1)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_summary_scatter(all_data_for_group, title, output_path):
    """生成汇总散点图（某场景某声源模式下所有频点叠加）"""
    all_delta = []
    all_freq = []
    all_spl = []

    for item in all_data_for_group:
        df = item['df']
        delta = df['delta_db'].dropna().astype(float).values
        freq = item['freq']
        all_delta.extend(delta)
        all_freq.extend([freq] * len(delta))

    all_delta = np.array(all_delta)
    all_freq = np.array(all_freq)
    x = np.arange(1, len(all_delta) + 1)

    fig, ax = plt.subplots(figsize=(14, 5))

    colors = []
    for d in all_delta:
        if d > 3:
            colors.append('#FF4444')
        elif d < -3:
            colors.append('#4444FF')
        else:
            colors.append('#44BB44')

    ax.scatter(x, all_delta, c=colors, s=15, alpha=0.6, edgecolors='none')

    # ±3dB红框
    rect = mpatches.FancyBboxPatch(
        (0, -3), len(all_delta) + 1, 6,
        boxstyle="round,pad=0",
        linewidth=2.5, edgecolor='red', facecolor='none',
        linestyle='--', zorder=5
    )
    ax.add_patch(rect)

    ax.axhline(y=3, color='red', linestyle='--', linewidth=1, alpha=0.7)
    ax.axhline(y=-3, color='red', linestyle='--', linewidth=1, alpha=0.7)
    ax.axhline(y=0, color='gray', linestyle='-', linewidth=0.8, alpha=0.5)

    total = len(all_delta)
    within_3 = ((all_delta >= -3) & (all_delta <= 3)).sum()
    above_3 = (all_delta > 3).sum()
    below_3 = (all_delta < -3).sum()

    ax.set_title(f'{title}\nΔdB(self-EASE5) 全频点汇总散点分布 (±3dB内占比: {within_3/total*100:.1f}%)',
                 fontsize=12, fontweight='bold')
    ax.set_xlabel('测点序号（全频点叠加）', fontsize=10)
    ax.set_ylabel('ΔdB (self - EASE5)', fontsize=10)

    from matplotlib.lines import Line2D
    legend_elements = [
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#44BB44', markersize=8,
               label=f'±3dB内 ({within_3}个, {within_3/total*100:.1f}%)'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#FF4444', markersize=8,
               label=f'>+3dB ({above_3}个, {above_3/total*100:.1f}%)'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#4444FF', markersize=8,
               label=f'<-3dB ({below_3}个, {below_3/total*100:.1f}%)'),
        Line2D([0], [0], color='red', linestyle='--', linewidth=2, label='±3dB容差线'),
    ]
    ax.legend(handles=legend_elements, loc='upper right', fontsize=8)

    ax.set_xlim(0, len(all_delta) + 1)
    y_range = max(abs(all_delta.min()), abs(all_delta.max()), 4)
    ax.set_ylim(-y_range - 1, y_range + 1)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()

    return {
        'total': total,
        'within_3_count': int(within_3),
        'within_3_pct': within_3/total*100,
        'above_3_count': int(above_3),
        'above_3_pct': above_3/total*100,
        'below_3_count': int(below_3),
        'below_3_pct': below_3/total*100,
        'mean': float(all_delta.mean()),
        'mae': float(np.abs(all_delta).mean()),
        'rmse': float(np.sqrt((all_delta**2).mean())),
        'std': float(all_delta.std()),
        'min': float(all_delta.min()),
        'max': float(all_delta.max()),
        'median': float(np.median(all_delta)),
        'p95_abs': float(np.percentile(np.abs(all_delta), 95)),
    }


def generate_freq_comparison_bar(all_data_for_group, title, output_path):
    """生成各频点±3dB占比柱状图"""
    freqs = []
    within_pcts = []
    above_pcts = []
    below_pcts = []

    for item in sorted(all_data_for_group, key=lambda x: FREQ_ORDER.index(x['freq']) if x['freq'] in FREQ_ORDER else 99):
        stats = compute_statistics(item['df'])
        if stats is None:
            continue
        freqs.append(f"{item['freq']}Hz")
        within_pcts.append(stats['within_3_pct'])
        above_pcts.append(stats['above_3_pct'])
        below_pcts.append(stats['below_3_pct'])

    x = np.arange(len(freqs))
    width = 0.6

    fig, ax = plt.subplots(figsize=(10, 5))
    bars1 = ax.bar(x, within_pcts, width, label='±3dB内', color='#44BB44', alpha=0.8)
    bars2 = ax.bar(x, above_pcts, width, bottom=within_pcts, label='>+3dB', color='#FF4444', alpha=0.8)
    bars3 = ax.bar(x, below_pcts, width, bottom=[w+a for w,a in zip(within_pcts, above_pcts)],
                   label='<-3dB', color='#4444FF', alpha=0.8)

    # 在柱子上标注百分比
    for i, (w, a, b) in enumerate(zip(within_pcts, above_pcts, below_pcts)):
        ax.text(i, w/2, f'{w:.1f}%', ha='center', va='center', fontsize=9, fontweight='bold')
        if a > 3:
            ax.text(i, w + a/2, f'{a:.1f}%', ha='center', va='center', fontsize=8, color='white')
        if b > 3:
            ax.text(i, w + a + b/2, f'{b:.1f}%', ha='center', va='center', fontsize=8, color='white')

    ax.set_title(title, fontsize=12, fontweight='bold')
    ax.set_xlabel('频点', fontsize=10)
    ax.set_ylabel('占比 (%)', fontsize=10)
    ax.set_xticks(x)
    ax.set_xticklabels(freqs)
    ax.legend(loc='upper right')
    ax.set_ylim(0, 110)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


def parse_source_coord(source_str):
    """解析声源坐标字符串"""
    if source_str is None:
        return None
    # 格式: 位置(m)：x:2.13，y:3.89，z:2.56\n角度(°)：水平:0，垂直:0，旋转:0
    coord_match = re.search(r'x:([\d.]+).*?y:([\d.]+).*?z:([\d.]+)', source_str)
    angle_match = re.search(r'水平:([\d.]+).*?垂直:([\d.]+).*?旋转:([\d.]+)', source_str)
    result = {}
    if coord_match:
        result['x'] = coord_match.group(1)
        result['y'] = coord_match.group(1)
        result['z'] = coord_match.group(3)
        # 修正y
        result['y'] = coord_match.group(2)
    if angle_match:
        result['水平'] = angle_match.group(1)
        result['垂直'] = angle_match.group(2)
        result['旋转'] = angle_match.group(3)
    return result if result else None


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_style(table):
    """给表格添加基本样式"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def generate_report(all_data):
    """生成Word测试报告"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('自研声场仿真 vs EASE5\n直达声压对比测试报告')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（基于RoomMapping求解结果）')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'报告日期：2026年5月8日')
    run.font.size = Pt(12)
    run = info.add_run(f'\n测试场景：中型教室1 / 会议室 / 超大型教室')
    run.font.size = Pt(12)
    run = info.add_run(f'\n声源模式：单声源 / 双声源')
    run.font.size = Pt(12)
    run = info.add_run(f'\n频段范围：125Hz ~ 8000Hz（1/3倍频程，A加权）')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ===== 目录 =====
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 测试概述',
        '2. 总体结论',
        '3. 中型教室1 - 直达声压对比',
        '   3.1 单声源',
        '   3.2 双声源',
        '4. 会议室 - 直达声压对比',
        '   4.1 单声源',
        '   4.2 双声源',
        '5. 超大型教室 - 直达声压对比',
        '   5.1 单声源',
        '   5.2 双声源',
        '6. 偏差超限测点明细',
        '7. 附录',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== 1. 测试概述 =====
    doc.add_heading('1. 测试概述', level=1)

    doc.add_heading('1.1 测试目的', level=2)
    doc.add_paragraph(
        '本报告旨在对比自研声场仿真系统与行业标杆EASE5在相同条件下的听音面直达声压计算结果，'
        '评估自研系统与EASE5之间的一致性水平，识别偏差较大的场景和频段，为后续算法优化提供数据支撑。'
    )

    doc.add_heading('1.2 测试方法', level=2)
    doc.add_paragraph(
        '在相同房间尺寸、声源参数和听音面网格条件下，分别使用自研系统和EASE5计算各测点的直达声压级（SPL），'
        '然后通过空间最近点位匹配（match_dist_xyz最小），对比同一位置的自研SPL与EASE5 SPL的差值 '
        'delta_db = self_spl - ease5_spl。以±3dB作为容差判定标准。'
    )

    doc.add_heading('1.3 测试场景', level=2)
    # 场景表格
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['场景', '声源模式', '频点数', '文件数']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '003366')
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    scenes_in_data = sorted(set(d['scene_display'] for d in all_data),
                            key=lambda x: ['中型教室1', '会议室', '超大型教室'].index(x) if x in ['中型教室1', '会议室', '超大型教室'] else 99)
    for i, scene in enumerate(scenes_in_data):
        scene_data = [d for d in all_data if d['scene_display'] == scene]
        row = table.rows[i + 1]
        row.cells[0].text = scene
        modes = set(d['source_display'] for d in scene_data)
        row.cells[1].text = '、'.join(sorted(modes))
        row.cells[2].text = str(len(set(d['freq'] for d in scene_data)))
        row.cells[3].text = str(len(scene_data))
    add_table_style(table)

    doc.add_heading('1.4 评判标准', level=2)
    p = doc.add_paragraph()
    run = p.add_run('核心指标：delta_db(self - EASE5)')
    run.font.bold = True
    doc.add_paragraph('• ±3dB以内：合格，自研与EASE5一致性良好')
    doc.add_paragraph('• >+3dB：自研声压偏高，需关注')
    doc.add_paragraph('• <-3dB：自研声压偏低，需关注')

    doc.add_page_break()

    # ===== 2. 总体结论 =====
    doc.add_heading('2. 总体结论', level=1)

    # 汇总所有数据
    all_delta = []
    for item in all_data:
        delta = item['df']['delta_db'].dropna().astype(float).values
        all_delta.extend(delta)
    all_delta = np.array(all_delta)
    total_all = len(all_delta)
    within_3_all = ((all_delta >= -3) & (all_delta <= 3)).sum()
    above_3_all = (all_delta > 3).sum()
    below_3_all = (all_delta < -3).sum()

    # 总体结论表
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    summary_data = [
        ('总测点数', f'{total_all}'),
        ('±3dB内占比', f'{within_3_all/total_all*100:.1f}%（{within_3_all}个）'),
        ('>+3dB占比', f'{above_3_all/total_all*100:.1f}%（{above_3_all}个）'),
        ('<-3dB占比', f'{below_3_all/total_all*100:.1f}%（{below_3_all}个）'),
        ('平均偏差', f'{all_delta.mean():.2f} dB'),
        ('平均绝对误差(MAE)', f'{np.abs(all_delta).mean():.2f} dB'),
        ('均方根误差(RMSE)', f'{np.sqrt((all_delta**2).mean()):.2f} dB'),
        ('P95绝对偏差', f'{np.percentile(np.abs(all_delta), 95):.2f} dB'),
    ]
    for i, (k, v) in enumerate(summary_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
    add_table_style(table)

    doc.add_paragraph()
    # 按场景+声源维度汇总
    doc.add_heading('2.1 各场景各声源模式汇总', level=2)

    summary_table = doc.add_table(rows=1, cols=8)
    summary_table.style = 'Table Grid'
    s_headers = ['场景', '声源模式', '测点数', '±3dB内占比', '>+3dB占比', '<-3dB占比', 'MAE(dB)', 'RMSE(dB)']
    for i, h in enumerate(s_headers):
        summary_table.rows[0].cells[i].text = h
        set_cell_shading(summary_table.rows[0].cells[i], '003366')
        for p in summary_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    # 生成各分组汇总图
    group_stats = {}
    for scene in scenes_in_data:
        for mode in ['单声源', '双声源']:
            group_data = [d for d in all_data if d['scene_display'] == scene and d['source_display'] == mode]
            if not group_data:
                continue
            group_delta = []
            for item in group_data:
                delta = item['df']['delta_db'].dropna().astype(float).values
                group_delta.extend(delta)
            group_delta = np.array(group_delta)
            if len(group_delta) == 0:
                continue

            total = len(group_delta)
            within_3 = ((group_delta >= -3) & (group_delta <= 3)).sum()
            above_3 = (group_delta > 3).sum()
            below_3 = (group_delta < -3).sum()

            gstats = {
                'total': total,
                'within_3_pct': within_3/total*100,
                'above_3_pct': above_3/total*100,
                'below_3_pct': below_3/total*100,
                'mae': np.abs(group_delta).mean(),
                'rmse': np.sqrt((group_delta**2).mean()),
                'mean': group_delta.mean(),
            }
            group_stats[(scene, mode)] = gstats

            row = summary_table.add_row()
            row.cells[0].text = scene
            row.cells[1].text = mode
            row.cells[2].text = str(total)
            row.cells[3].text = f'{gstats["within_3_pct"]:.1f}%'
            row.cells[4].text = f'{gstats["above_3_pct"]:.1f}%'
            row.cells[5].text = f'{gstats["below_3_pct"]:.1f}%'
            row.cells[6].text = f'{gstats["mae"]:.2f}'
            row.cells[7].text = f'{gstats["rmse"]:.2f}'

            # 高亮±3dB占比低的
            if gstats['within_3_pct'] < 80:
                set_cell_shading(row.cells[3], 'FFCCCC')
            elif gstats['within_3_pct'] >= 95:
                set_cell_shading(row.cells[3], 'CCFFCC')

    add_table_style(summary_table)

    doc.add_page_break()

    # ===== 3-5. 各场景详细分析 =====
    section_num = 3
    for scene in scenes_in_data:
        doc.add_heading(f'{section_num}. {scene} - 直达声压对比', level=1)

        # 收集声源坐标信息
        scene_data = [d for d in all_data if d['scene_display'] == scene]
        self_sources = set()
        ease5_sources = set()
        for item in scene_data:
            if item['sources'][0]:
                self_sources.add(item['sources'][0])
            if item['sources'][1]:
                ease5_sources.add(item['sources'][1])

        # 声源坐标信息
        doc.add_heading(f'{section_num}.0 声源坐标', level=2)
        src_table = doc.add_table(rows=1, cols=2)
        src_table.style = 'Table Grid'
        src_table.rows[0].cells[0].text = '自研声源坐标'
        src_table.rows[0].cells[1].text = 'EASE5声源坐标'
        set_cell_shading(src_table.rows[0].cells[0], '003366')
        set_cell_shading(src_table.rows[0].cells[1], '003366')
        for p in src_table.rows[0].cells[0].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
        for p in src_table.rows[0].cells[1].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

        row = src_table.add_row()
        # 去重显示
        self_src_text = '\n'.join(sorted(self_sources)) if self_sources else '无数据'
        ease5_src_text = '\n'.join(sorted(ease5_sources)) if ease5_sources else '无数据'
        row.cells[0].text = self_src_text
        row.cells[1].text = ease5_src_text
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(8)

        sub_num = 1
        for mode in ['单声源', '双声源']:
            doc.add_heading(f'{section_num}.{sub_num} {mode}', level=2)

            group_data = [d for d in all_data if d['scene_display'] == scene and d['source_display'] == mode]
            if not group_data:
                doc.add_paragraph('暂无数据')
                sub_num += 1
                continue

            # 汇总散点图
            summary_chart_path = os.path.join(CHART_DIR, f'summary_{scene}_{mode}.png')
            gstats = generate_summary_scatter(group_data, f'{scene} - {mode} - 全频点汇总', summary_chart_path)
            doc.add_picture(summary_chart_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 频点占比柱状图
            bar_chart_path = os.path.join(CHART_DIR, f'bar_{scene}_{mode}.png')
            generate_freq_comparison_bar(group_data, f'{scene} - {mode} - 各频点±3dB占比', bar_chart_path)
            doc.add_picture(bar_chart_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 各频点详细表格
            doc.add_heading(f'{section_num}.{sub_num}.1 各频点统计明细', level=3)

            freq_table = doc.add_table(rows=1, cols=10)
            freq_table.style = 'Table Grid'
            f_headers = ['频点', '测点数', '±3dB内占比', '>+3dB占比', '<-3dB占比',
                         '均值(dB)', 'MAE(dB)', 'RMSE(dB)', 'P95|Δ|(dB)', '最大|Δ|(dB)']
            for i, h in enumerate(f_headers):
                freq_table.rows[0].cells[i].text = h
                set_cell_shading(freq_table.rows[0].cells[i], '003366')
                for p in freq_table.rows[0].cells[i].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.bold = True
                        r.font.size = Pt(8)

            for item in sorted(group_data, key=lambda x: FREQ_ORDER.index(x['freq']) if x['freq'] in FREQ_ORDER else 99):
                stats = compute_statistics(item['df'])
                if stats is None:
                    continue
                row = freq_table.add_row()
                row.cells[0].text = f"{item['freq']}Hz"
                row.cells[1].text = str(stats['total'])
                row.cells[2].text = f'{stats["within_3_pct"]:.1f}%'
                row.cells[3].text = f'{stats["above_3_pct"]:.1f}%'
                row.cells[4].text = f'{stats["below_3_pct"]:.1f}%'
                row.cells[5].text = f'{stats["mean"]:.2f}'
                row.cells[6].text = f'{stats["mae"]:.2f}'
                row.cells[7].text = f'{stats["rmse"]:.2f}'
                row.cells[8].text = f'{stats["p95_abs"]:.2f}'
                row.cells[9].text = f'{stats["max_abs_db"]:.2f}' if 'max_abs_db' in stats else f'{stats["max"]:.2f}'

                # 高亮
                if stats['within_3_pct'] < 80:
                    set_cell_shading(row.cells[2], 'FFCCCC')
                elif stats['within_3_pct'] >= 95:
                    set_cell_shading(row.cells[2], 'CCFFCC')

            add_table_style(freq_table)

            # 各频点散点图
            doc.add_heading(f'{section_num}.{sub_num}.2 各频点散点图', level=3)
            for item in sorted(group_data, key=lambda x: FREQ_ORDER.index(x['freq']) if x['freq'] in FREQ_ORDER else 99):
                chart_path = os.path.join(CHART_DIR, f'scatter_{scene}_{mode}_{item["freq"]}Hz.png')
                generate_scatter_plot(item, chart_path)
                doc.add_picture(chart_path, width=Inches(5.8))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            sub_num += 1

        doc.add_page_break()
        section_num += 1

    # ===== 6. 偏差超限测点明细 =====
    doc.add_heading('6. 偏差超限测点明细', level=1)
    doc.add_paragraph('以下列出所有 |delta_db| > 3dB 的测点详细信息，供开发团队定位优化。')

    for scene in scenes_in_data:
        for mode in ['单声源', '双声源']:
            group_data = [d for d in all_data if d['scene_display'] == scene and d['source_display'] == mode]
            if not group_data:
                continue

            doc.add_heading(f'{scene} - {mode}', level=2)

            outlier_rows = []
            for item in sorted(group_data, key=lambda x: FREQ_ORDER.index(x['freq']) if x['freq'] in FREQ_ORDER else 99):
                df = item['df']
                for _, row in df.iterrows():
                    try:
                        delta = float(row['delta_db'])
                    except (ValueError, TypeError):
                        continue
                    if abs(delta) > 3:
                        outlier_rows.append({
                            'freq': item['freq'],
                            'ease5_x': row.get('ease5_x', ''),
                            'ease5_y': row.get('ease5_y', ''),
                            'ease5_z': row.get('ease5_z', ''),
                            'self_x': row.get('self_x', ''),
                            'self_y': row.get('self_y', ''),
                            'self_z': row.get('self_z', ''),
                            'ease5_spl': row.get('ease5_spl', ''),
                            'self_spl': row.get('self_spl', ''),
                            'delta_db': delta,
                            'match_dist_xyz': row.get('match_dist_xyz', ''),
                            'ease5_seq': row.get('ease5_seq', ''),
                            'self_seq': row.get('self_seq', ''),
                            'ease5_row': row.get('ease5_row', ''),
                            'ease5_col': row.get('ease5_col', ''),
                            'self_row': row.get('self_row', ''),
                            'self_col': row.get('self_col', ''),
                        })

            if not outlier_rows:
                doc.add_paragraph('无超限测点，全部在±3dB以内。')
                continue

            doc.add_paragraph(f'共 {len(outlier_rows)} 个测点偏差超过±3dB：')

            # 表格
            o_table = doc.add_table(rows=1, cols=10)
            o_table.style = 'Table Grid'
            o_headers = ['频点', 'EASE5坐标(x,y,z)', '自研坐标(x,y,z)', 'EASE5 SPL', '自研 SPL',
                         'ΔdB', '匹配距离', 'EASE5序号', '自研序号(行,列)', '偏差方向']
            for i, h in enumerate(o_headers):
                o_table.rows[0].cells[i].text = h
                set_cell_shading(o_table.rows[0].cells[i], '003366')
                for p in o_table.rows[0].cells[i].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.bold = True
                        r.font.size = Pt(7)

            # 按偏差绝对值排序，取前50条
            outlier_rows.sort(key=lambda x: abs(x['delta_db']), reverse=True)
            for orow in outlier_rows[:50]:
                row = o_table.add_row()
                row.cells[0].text = f"{orow['freq']}Hz"
                row.cells[1].text = f"({orow['ease5_x']},{orow['ease5_y']},{orow['ease5_z']})"
                row.cells[2].text = f"({orow['self_x']},{orow['self_y']},{orow['self_z']})"
                try:
                    row.cells[3].text = f"{float(orow['ease5_spl']):.2f}"
                except:
                    row.cells[3].text = str(orow['ease5_spl'])
                try:
                    row.cells[4].text = f"{float(orow['self_spl']):.2f}"
                except:
                    row.cells[4].text = str(orow['self_spl'])
                row.cells[5].text = f"{orow['delta_db']:.2f}"
                try:
                    row.cells[6].text = f"{float(orow['match_dist_xyz']):.4f}"
                except:
                    row.cells[6].text = str(orow['match_dist_xyz'])
                row.cells[7].text = str(orow['ease5_seq'])
                row.cells[8].text = f"({orow['self_row']},{orow['self_col']})"
                row.cells[9].text = '自研偏高' if orow['delta_db'] > 0 else '自研偏低'

                # 偏差方向着色
                if orow['delta_db'] > 3:
                    set_cell_shading(row.cells[9], 'FFCCCC')
                elif orow['delta_db'] < -3:
                    set_cell_shading(row.cells[9], 'CCCCFF')

            if len(outlier_rows) > 50:
                doc.add_paragraph(f'... 仅显示前50条，共 {len(outlier_rows)} 条超限记录')

            add_table_style(o_table)

    # ===== 7. 附录 =====
    doc.add_page_break()
    doc.add_heading('7. 附录', level=1)

    doc.add_heading('7.1 术语说明', level=2)
    terms = [
        ('delta_db(self-EASE5)', '自研系统计算的声压级减去EASE5计算的声压级，单位dB'),
        ('SPL', 'Sound Pressure Level，声压级'),
        ('MAE', 'Mean Absolute Error，平均绝对误差'),
        ('RMSE', 'Root Mean Square Error，均方根误差'),
        ('match_dist_xyz', '自研测点与EASE5测点的空间匹配距离（米）'),
        ('±3dB容差', '工程上一般认为±3dB以内为可接受偏差'),
    ]
    t_table = doc.add_table(rows=len(terms)+1, cols=2)
    t_table.style = 'Table Grid'
    t_table.rows[0].cells[0].text = '术语'
    t_table.rows[0].cells[1].text = '说明'
    set_cell_shading(t_table.rows[0].cells[0], '003366')
    set_cell_shading(t_table.rows[0].cells[1], '003366')
    for p in t_table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True
    for p in t_table.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True
    for i, (term, desc) in enumerate(terms):
        t_table.rows[i+1].cells[0].text = term
        t_table.rows[i+1].cells[1].text = desc
    add_table_style(t_table)

    doc.add_heading('7.2 数据来源', level=2)
    doc.add_paragraph('所有数据来源于RoomMapping求解结果，自研系统与EASE5使用相同房间参数和声源配置。')
    doc.add_paragraph(f'数据文件路径：{BASE_DIR}')

    # 保存
    doc.save(REPORT_FILE)
    print(f"\n报告已生成: {REPORT_FILE}")


if __name__ == '__main__':
    print("=" * 60)
    print("自研 vs EASE5 直达声压对比测试报告生成")
    print("=" * 60)

    all_data = collect_all_data()
    print(f"\n共读取 {len(all_data)} 个有效数据集")

    if not all_data:
        print("无有效数据，退出")
        exit(1)

    generate_report(all_data)
    print("\n完成！")
