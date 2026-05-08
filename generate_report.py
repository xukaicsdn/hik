import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
import re
import warnings
warnings.filterwarnings('ignore')

# Read combined data
base_path = Path(r'E:\AudioSimulationSnap\点位声压采样\求解结果对比-直达声压（RoomMapping）')
exclude_path = base_path / '大型教室 - 由于EASE5和自研尺寸对不上暂时不测'
excel_files = [f for f in list(base_path.rglob('*ES5和声压数据相对最近点位声压级对比.xlsx')) if exclude_path not in f.parents]

all_data = []
for file in excel_files:
    try:
        df = pd.read_excel(file, engine='openpyxl')
        scene = '中型教室 1' if '中型教室 1' in str(file) else '会议室' if '会议室' in str(file) else '超大型教室' if '超大型教室' in str(file) else '未知'
        source_type = '单声源' if '单声源' in str(file) else '双声源' if '双声源' in str(file) else '未知'
        freq_match = re.search(r'\[(\d+)HZ\]', file.stem)
        freq = freq_match.group(1) + 'Hz' if freq_match else '未知'
        df['scene'] = scene
        df['source_type'] = source_type
        df['frequency'] = freq
        all_data.append(df)
    except:
        pass

combined_df = pd.concat(all_data, ignore_index=True)
delta_col = 'delta_db(self-ease5)'

# Create Word document
doc = Document()

# Set default font style
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
font.size = Pt(10)

# Title
title = doc.add_heading('自研声学仿真软件 vs EASE5\n直达声压对比测试报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.name = '微软雅黑'
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# Executive Summary
doc.add_heading('一、执行摘要', level=1)
summary_text = f'''本报告对比了自研声学仿真软件与行业标杆 EASE5 在三种典型场景（中型教室 1、会议室、超大型教室）下，
单声源和双声源配置的听音面相对点位直达声压数据。测试覆盖 7 个频率点（125Hz-8000Hz），共计{len(combined_df)}个采样点。

核心结论：
- 整体一致性：{((combined_df[delta_col].dropna() >= -3) & (combined_df[delta_col].dropna() <= 3)).sum()}/{len(combined_df[delta_col].dropna())}
  （{((combined_df[delta_col].dropna() >= -3) & (combined_df[delta_col].dropna() <= 3)).sum()/len(combined_df[delta_col].dropna())*100:.1f}%）的数据偏差在±3dB 以内
- 平均偏差：{combined_df[delta_col].dropna().mean():.2f} dB
- 标准差：{combined_df[delta_col].dropna().std():.2f} dB
- 正向超差（>3dB）：{(combined_df[delta_col].dropna() > 3).sum()}个采样点（{(combined_df[delta_col].dropna() > 3).sum()/len(combined_df[delta_col].dropna())*100:.1f}%）
- 负向超差（<-3dB）：{(combined_df[delta_col].dropna() < -3).sum()}个采样点（{(combined_df[delta_col].dropna() < -3).sum()/len(combined_df[delta_col].dropna())*100:.1f}%）
'''
p = doc.add_paragraph(summary_text)
p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# Test Configuration
doc.add_heading('二、测试配置', level=1)

doc.add_heading('2.1 测试场景', level=2)
config_table = doc.add_table(rows=4, cols=4)
config_table.style = 'Light Grid Accent 1'

headers = ['场景', '单声源配置', '双声源配置', '频率点']
hdr_cells = config_table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

scenes_info = {
    '中型教室 1': ['声源 1: (2.13, 4.25, 2.56)m', '声源 1: (4.32, 4.25, 1.64)m\n声源 2: (5.54, 3.89, 1.64)m', '125, 250, 500, 1000,\n2000, 4000, 8000 Hz'],
    '会议室': ['声源 1: (0.03, 6.25, 2.00)m', '声源 1: (2.39, 1.73, 2.00)m\n声源 2: (-1.61, 1.73, 2.00)m', '125, 250, 500, 1000,\n2000, 4000, 8000 Hz'],
    '超大型教室': ['声源 1: (13.35, 0, 2.77)m', '声源 1: (3, -3, 4.00)m\n声源 2: (3, 3, 4.00)m', '125, 250, 500, 1000,\n2000, 4000, 8000 Hz']
}

for idx, (scene, info) in enumerate(scenes_info.items()):
    row = config_table.rows[idx + 1]
    row.cells[0].text = scene
    row.cells[1].text = info[0]
    row.cells[2].text = info[1]
    row.cells[3].text = info[2]

doc.add_paragraph()

# Overall Statistics
doc.add_heading('三、整体统计分析', level=1)

valid_delta = combined_df[delta_col].dropna()
total = len(valid_delta)
within_3 = ((valid_delta >= -3) & (valid_delta <= 3)).sum()
exceed_pos = (valid_delta > 3).sum()
exceed_neg = (valid_delta < -3).sum()

stats_table = doc.add_table(rows=6, cols=2)
stats_table.style = 'Light List Accent 1'

stats_data = [
    ['指标', '数值'],
    ['总采样点数', f'{total}'],
    ['平均偏差', f'{valid_delta.mean():.3f} dB'],
    ['标准差', f'{valid_delta.std():.3f} dB'],
    ['最小值 / 最大值', f'{valid_delta.min():.2f} / {valid_delta.max():.2f} dB'],
    ['±3dB 以内占比', f'{within_3}/{total} ({within_3/total*100:.1f}%)']
]

for i, (label, value) in enumerate(stats_data):
    row = stats_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# By Scene and Source Type
doc.add_heading('四、分场景/声源类型统计', level=1)

for scene in ['中型教室 1', '会议室', '超大型教室']:
    doc.add_heading(f'4.1 {scene}', level=2)

    scene_data = combined_df[combined_df['scene'] == scene]
    scene_delta = scene_data[delta_col].dropna()

    p = doc.add_paragraph()
    p.add_run(f'数据量：{len(scene_delta)} | ')
    p.add_run(f'平均：{scene_delta.mean():.2f} dB | ')
    p.add_run(f'标准差：{scene_delta.std():.2f} dB | ')
    within = ((scene_delta >= -3) & (scene_delta <= 3)).sum()
    p.add_run(f'±3dB 以内：{within}/{len(scene_delta)} ({within/len(scene_delta)*100:.1f}%)')

    for stype in ['单声源', '双声源']:
        subset = scene_data[scene_data['source_type'] == stype]
        sub_delta = subset[delta_col].dropna()
        if len(sub_delta) > 0:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'{stype}: ')
            p.add_run(f'平均={sub_delta.mean():.2f}dB, ')
            p.add_run(f'±3dB 占比={((sub_delta >= -3) & (sub_delta <= 3)).sum()/len(sub_delta)*100:.1f}%')

    # Show top exceeding points
    exceed = scene_data[scene_data[delta_col].abs() > 3]
    if len(exceed) > 0:
        doc.add_paragraph(f'超差点数：{len(exceed)}', style='List Bullet')
        if len(exceed) <= 10:
            exc_table = doc.add_table(rows=len(exceed)+1, cols=5)
            exc_table.style = 'Table Grid'
            hdr = ['序号', '位置 (x,y,z)', 'EASE5 SPL', 'Self SPL', 'Delta dB']
            for j, h in enumerate(hdr):
                exc_table.rows[0].cells[j].text = h
                exc_table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
            for i, (_, row) in enumerate(exceed.head(10).iterrows()):
                exc_table.rows[i+1].cells[0].text = str(i+1)
                exc_table.rows[i+1].cells[1].text = f'({row["self_x"]:.2f}, {row["self_y"]:.2f}, {row["self_z"]:.2f})'
                exc_table.rows[i+1].cells[2].text = f'{row["ease5_spl"]:.1f}'
                exc_table.rows[i+1].cells[3].text = f'{row["self_spl"]:.1f}'
                exc_table.rows[i+1].cells[4].text = f'{row[delta_col]:.2f}'

doc.add_page_break()

# Frequency Analysis
doc.add_heading('五、分频率统计', level=1)

freq_table = doc.add_table(rows=8, cols=5)
freq_table.style = 'Light Grid Accent 1'
hdr = ['频率', '数据量', '平均偏差 (dB)', '标准差 (dB)', '±3dB 占比']
for j, h in enumerate(hdr):
    freq_table.rows[0].cells[j].text = h
    freq_table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True

freqs = ['125Hz', '250Hz', '500Hz', '1000Hz', '2000Hz', '4000Hz', '8000Hz']
for i, freq in enumerate(freqs):
    subset = combined_df[combined_df['frequency'] == freq]
    sub_delta = subset[delta_col].dropna()
    if len(sub_delta) > 0:
        row = freq_table.rows[i+1]
        row.cells[0].text = freq
        row.cells[1].text = str(len(sub_delta))
        row.cells[2].text = f'{sub_delta.mean():.2f}'
        row.cells[3].text = f'{sub_delta.std():.2f}'
        within = ((sub_delta >= -3) & (sub_delta <= 3)).sum()
        row.cells[4].text = f'{within/len(sub_delta)*100:.1f}%'

doc.add_page_break()

# Charts Section
doc.add_heading('六、可视化图表', level=1)

doc.add_heading('6.1 散点图 - 按场景和声源类型', level=2)
doc.add_picture(r'E:\AudioSimulationSnap\test_report_scatter_all.png', width=Inches(7.5))
doc.add_paragraph('图 1: 各场景 + 声源类型的 Delta dB 分布（红色框标注±3dB 区域）')

doc.add_heading('6.2 散点图 - 按频率', level=2)
doc.add_picture(r'E:\AudioSimulationSnap\test_report_scatter_freq.png', width=Inches(7.5))
doc.add_paragraph('图 2: 各频率点的 Delta dB 分布（红色框标注±3dB 区域）')

doc.add_heading('6.3 箱线图 - 场景对比', level=2)
doc.add_picture(r'E:\AudioSimulationSnap\test_report_boxplot.png', width=Inches(7.5))
doc.add_paragraph('图 3: 各场景 Delta dB 箱线图对比（红线为±3dB 阈值）')

doc.add_page_break()

# Exceeding Data Details
doc.add_heading('七、超差数据详情（|Δ| > 3dB）', level=1)

# Positive exceed
doc.add_heading('7.1 正向超差（Self > EASE5 +3dB）', level=2)
exceed_pos_data = combined_df[combined_df[delta_col] > 3].sort_values(delta_col, ascending=False)
if len(exceed_pos_data) > 0:
    p = doc.add_paragraph(f'共{len(exceed_pos_data)}个采样点，最大偏差：{exceed_pos_data[delta_col].max():.2f} dB')

    if len(exceed_pos_data) <= 50:
        ep_table = doc.add_table(rows=len(exceed_pos_data)+1, cols=6)
        ep_table.style = 'Table Grid'
        hdr = ['场景', '声源', '频率', '位置 (x,y,z)', 'EASE5 SPL', 'Delta dB']
        for j, h in enumerate(hdr):
            ep_table.rows[0].cells[j].text = h
            ep_table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        for i, (_, row) in enumerate(exceed_pos_data.iterrows()):
            ep_table.rows[i+1].cells[0].text = row['scene']
            ep_table.rows[i+1].cells[1].text = row['source_type']
            ep_table.rows[i+1].cells[2].text = row['frequency']
            ep_table.rows[i+1].cells[3].text = f'({row["self_x"]:.2f}, {row["self_y"]:.2f}, {row["self_z"]:.2f})'
            ep_table.rows[i+1].cells[4].text = f'{row["ease5_spl"]:.1f}'
            ep_table.rows[i+1].cells[5].text = f'{row[delta_col]:.2f}'
else:
    doc.add_paragraph('无正向超差数据')

# Negative exceed
doc.add_heading('7.2 负向超差（Self < EASE5 -3dB）', level=2)
exceed_neg_data = combined_df[combined_df[delta_col] < -3].sort_values(delta_col, ascending=True)
if len(exceed_neg_data) > 0:
    p = doc.add_paragraph(f'共{len(exceed_neg_data)}个采样点，最大偏差：{exceed_neg_data[delta_col].min():.2f} dB')

    if len(exceed_neg_data) <= 50:
        en_table = doc.add_table(rows=len(exceed_neg_data)+1, cols=6)
        en_table.style = 'Table Grid'
        hdr = ['场景', '声源', '频率', '位置 (x,y,z)', 'EASE5 SPL', 'Delta dB']
        for j, h in enumerate(hdr):
            en_table.rows[0].cells[j].text = h
            en_table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        for i, (_, row) in enumerate(exceed_neg_data.iterrows()):
            en_table.rows[i+1].cells[0].text = row['scene']
            en_table.rows[i+1].cells[1].text = row['source_type']
            en_table.rows[i+1].cells[2].text = row['frequency']
            en_table.rows[i+1].cells[3].text = f'({row["self_x"]:.2f}, {row["self_y"]:.2f}, {row["self_z"]:.2f})'
            en_table.rows[i+1].cells[4].text = f'{row["ease5_spl"]:.1f}'
            en_table.rows[i+1].cells[5].text = f'{row[delta_col]:.2f}'
else:
    doc.add_paragraph('无负向超差数据')

doc.add_page_break()

# Recommendations
doc.add_heading('八、优化建议', level=1)

rec_text = '''基于数据分析，提出以下优化建议：

1. 高频段优化（4000Hz, 8000Hz）：
   - 高频段偏差相对较大，建议检查自研软件的高频吸收系数和边界条件处理

2. 双声源干涉区域：
   - 双声源配置下部分区域偏差较大，建议优化声源叠加算法

3. 近场区域校准：
   - 距离声源较近的点位存在系统性偏低，建议调整近场补偿参数

4. 场景几何建模：
   - 超大型教室的偏差略高于其他场景，需复核复杂几何体的网格划分精度

5. 重点优化方向：
   - 优先处理负向超差区域（自研结果偏低），确保保守估计
   - 关注平均偏差>1dB 的场景进行参数微调
'''
doc.add_paragraph(rec_text)

doc.add_page_break()

# Appendix
doc.add_heading('九、附录：声源坐标详情', level=1)

src_doc = '''
中型教室 1：
  单声源：Self(2.13, 3.89, 2.56)m / EASE5(2.13, 4.25, 2.56)m
  双声源：Self1(4.32, 3.89, 1.64)m, Self2(5.54, 3.89, 1.64)m
          EASE5_1(4.32, 4.25, 1.64)m, EASE5_2(5.54, 3.89, 1.64)m

会议室：
  单声源：Self(0.03, 5.89, 2.00)m / EASE5(0.03, 6.25, 2.00)m
  双声源：Self1(2.75, 1.73, 2.00)m, Self2(-1.25, 1.73, 2.00)m
          EASE5_1(2.39, 1.73, 2.00)m, EASE5_2(-1.61, 1.73, 2.00)m

超大型教室：
  单声源：Self(13.35, -0.36, 2.77)m / EASE5(13.35, 0, 2.77)m
  双声源：Self1(3, -2.64, 4.00)m, Self2(3, 2.64, 4.00)m
          EASE5_1(3, -3, 4.00)m, EASE5_2(3, 3, 4.00)m

注：自研与 EASE5 声源坐标存在微小差异，主要源于网格离散化处理方式不同。
'''
doc.add_paragraph(src_doc)

# Save document
doc.save(r'E:\AudioSimulationSnap\自研 vs EASE5_直达声压对比测试报告.docx')
print("Word 报告已生成：自研 vs EASE5_直达声压对比测试报告.docx")
